#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成示例Word题库文档
帮助用户了解正确的格式规范
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    print("❌ 请先安装python-docx: pip install python-docx")
    exit(1)

doc = Document()

# 标题
title = doc.add_heading('刷题系统 - 示例题库', 0)
title.alignment = 1  # 居中

# 说明
doc.add_paragraph('本示例展示了正确的题目格式，请按照此格式编写题库。')
doc.add_paragraph('')

# 单选题
doc.add_heading('一、单选题', level=2)
doc.add_paragraph('1. 下列哪个是Python的Web框架？')
doc.add_paragraph('A. Django')
doc.add_paragraph('B. NumPy')
doc.add_paragraph('C. Pandas')
doc.add_paragraph('D. Matplotlib')
p = doc.add_paragraph('答案：A')
p.runs[0].bold = True

doc.add_paragraph('')

# 多选题
doc.add_heading('二、多选题', level=2)
doc.add_paragraph('2. 下列哪些是Python的Web框架？（多选）')
doc.add_paragraph('A. Django')
doc.add_paragraph('B. Flask')
doc.add_paragraph('C. NumPy')
doc.add_paragraph('D. Tornado')
p = doc.add_paragraph('答案：A,B,D')
p.runs[0].bold = True

doc.add_paragraph('')

# 判断题
doc.add_heading('三、判断题', level=2)
doc.add_paragraph('3. Python是一种编译型语言。')
p = doc.add_paragraph('答案：错')
p.runs[0].bold = True

doc.add_paragraph('4. Flask是一个轻量级的Python Web框架。')
p = doc.add_paragraph('答案：对')
p.runs[0].bold = True

doc.add_paragraph('')

# 填空题
doc.add_heading('四、填空题', level=2)
doc.add_paragraph('5. Python的Web框架____可以快速开发Web应用。')
p = doc.add_paragraph('答案：Django')
p.runs[0].bold = True

doc.add_paragraph('')

# 简答题
doc.add_heading('五、简答题', level=2)
doc.add_paragraph('6. 简述Django的MVT架构。')
p = doc.add_paragraph('答案：Model（模型）负责数据处理，View（视图）负责业务逻辑，Template（模板）负责页面展示。')
p.runs[0].bold = True

# 保存
doc.save('示例题库.docx')
print("✅ 示例题库已生成: 示例题库.docx")
print("📖 请打开查看格式规范")
